from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from web_novel_meta_infos import WebNovelMetaInfos
from novel import Novel
import time
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from docx2pdf import convert
import aspose.words as aw

class CreateBook:

    def getChapterContent(driver, className, classType, removeList = None):
        clType = ""
        match classType:
            case "id":
                clType = By.ID
            case "class":
                clType = By.CLASS_NAME
            case _:
                clType = ""
                print("Wrong classType!")
        
        chapter = driver.find_element(clType, className).text

        if removeList is not None:
            for remove in removeList:
                chapter = chapter.replace(remove, '')
        
        return chapter


    def createDoc(novel, url, removeList = None):
        """
        url needs to contain chapter-firstCh, otherwise it won't work
        """
        doc = docx.Document()
        doc = CreateBook.createTitelPages(doc, novel.imagePath, novel.titel, novel.author)

        driver = webdriver.Chrome()
        driver.get(url)
        time.sleep(5)
        counter = novel.firstCh
        for ch in range(novel.firstCh, novel.lastCh+1):
            time.sleep(5)
            header = doc.add_heading(f"Chapter {ch}", 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER

            chapter = None
            chapter = CreateBook.getChapterContent(driver, novel.className, novel.classType, removeList)
            doc.add_paragraph(chapter)
            print(f"collected chapter: {ch}")
            if ch == novel.lastCh:
                break
            """
            what follows is the navigation to the next page
            """
            doc.add_page_break()
            counter = counter + 1
            time.sleep(5)
            url = url.replace(f"chapter-{ch}", f"chapter-{counter}")
            driver.get(url)

        doc.save(f"C:/Users/mathi/OneDrive/Dokumente/ProgrammingProjects/webnovelscraper/public/novels/{novel.titel}.docx")
        driver.quit()

    def createTitelPages(doc, imagePath, titel, author):
        titelPageSection = doc.sections[-1]
        titelPageSection.bottom_margin = Inches(0)
        titelPageSection.left_margin = Inches(0)
        titelPageSection.right_margin = Inches(0)
        titelPageSection.top_margin = Inches(0)
        doc.add_picture(imagePath, width=docx.shared.Inches(8.5), height=docx.shared.Inches(11))

        chapterPagesSection = doc.sections[0]
        chapterPagesSection = doc.add_section(WD_SECTION.NEW_PAGE)
        chapterPagesSection.bottom_margin = Inches(1)
        chapterPagesSection.left_margin = Inches(1.25)
        chapterPagesSection.right_margin = Inches(1.25)
        chapterPagesSection.top_margin = Inches(1)
        bookTitel = doc.add_heading(titel, 0)
        bookTitel.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bookTitel = doc.add_heading("by " + author, 1)
        bookTitel.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
        return doc
    
    def convertWordToPdf(novel):
        print("*************")
        print("starting conversion to pdf")
        print("*************")
        convert(f"C:/Users/mathi/OneDrive/Dokumente/ProgrammingProjects/webnovelscraper/public/novels/{novel.titel}.docx", 
                f"C:/Users/mathi/OneDrive/Dokumente/ProgrammingProjects/webnovelscraper/public/novels/{novel.titel}.pdf")
        print("*************")
        print("conversion finished")
        print("*************")

    def testCreateDoc():
        # Trash of the Counts family
        n = WebNovelMetaInfos.getTitleAndDescription("https://www.novelupdates.com/series/trash-of-the-counts-family/")
        n.setChapterInfos(1, 3, "entry-content", "class")
        removeList = ["<< Previous Chapter | Index | Next Chapter >>"]
        CreateBook.createDoc(n, "https://eatapplepies.com/tcf-chapter-1/", removeList)
        CreateBook.convertWordToPdf(n)
        # Otome game no heroine
        n = WebNovelMetaInfos.getTitleAndDescription("https://www.novelupdates.com/series/otome-game-no-heroine-de-saikyou-survival/")
        n.setChapterInfos(1, 3, "entry-content", "class")
        removeList = ["Prev  Table of Content   Next", "Bagikan ini:", "TwitterFacebook", "Loading..."]
        CreateBook.createDoc(n, "https://bakapervert.wordpress.com/heroine-survival-vol-1-chapter-1/", removeList)
        CreateBook.convertWordToPdf(n)
        # omnisicient readers viewpoint ==== got hit wit cloudfare
        # n = WebNovelMetaInfos.getTitleAndDescription("https://www.novelupdates.com/series/omniscient-readers-viewpoint/")
        # CreateBook.createDoc(n, 1, 3, "https://novelusb.com/novel-book/omniscient-readers-viewpoint-novel/chapter-1", "chr-content", "id")
        # CreateBook.convertWordToPdf(n)
        

    def testGetChapterContent():
        driver = webdriver.Chrome()
        # Trash of the Counts family
        print("test: Trash of the Counts family")
        driver.get("https://eatapplepies.com/tcf-chapter-1/")
        time.sleep(10)
        removeList = ["<< Previous Chapter | Index | Next Chapter >>"]
        chapter = CreateBook.getChapterContent(driver, "entry-content", "class", removeList)
        print(chapter)
        # Otome game no heroine
        print("test: Otome game no heroine")
        time.sleep(10)
        removeList = ["Prev  Table of Content   Next", "Bagikan ini:", "TwitterFacebook", "Loading..."]
        driver.get("https://bakapervert.wordpress.com/heroine-survival-vol-1-chapter-1/")
        time.sleep(10)
        chapter = CreateBook.getChapterContent(driver, "entry-content", "class", removeList)
        print(chapter)
        # omnisicient readers viewpoint 
        print("test: omnisicient readers viewpoint")
        time.sleep(10)
        driver.get("https://novelusb.com/novel-book/omniscient-readers-viewpoint-novel/chapter-1")
        time.sleep(10)
        chapter = CreateBook.getChapterContent(driver, "chr-content", "id") 
        print(chapter)




        






